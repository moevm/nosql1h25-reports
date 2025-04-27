import docx
import os
from typing import Union, BinaryIO
from io import BytesIO

from src.diploma_processing.parsing_docx.docClasses import Doc, DocSection


class DocxParser:
    def __init__(self, file: str | Union[BinaryIO, BytesIO]):
        self.file = file
        self.docx_file = None
        self.doc = None

    def _check_file_exists(self):
        """
        Проверяет, существует ли файл и является ли он строкой.
        
        :raises ValueError: если путь к файлу не валиден
        :raises FileNotFoundError: если файл не существует
        """
        if not self.file or not isinstance(self.file, str):
            raise ValueError("Путь к файлу должен быть строкой и не пустым")
        try:
            if not os.path.exists(self.file) or not os.path.isfile(self.file):
                raise FileNotFoundError(f"Файл '{self.file}' не найден")
        except ImportError as e:
            raise("Возможная ошибка: ", str(e))

    def _check_file_format(self):
        """
        Проверяет, является ли файл форматом .docx.
        
        :raises ValueError: если файл имеет неверный формат
        """
        try:
            self.docx_file = docx.Document(self.file)
        except docx.opc.exceptions.PackageNotFoundError as e:
            raise ValueError(f"Файл '{self.file}' имеет неверный формат (.docx) или поврежден")
        except Exception as e:
            raise ValueError(f"Ошибка при открытии файла '{self.file}': {e}")
    
    def _load_file_from_buffer(self):
        """
        Предполагается, что на вход подан какой-то буффер или массив байтов.
        
        :raises ValueError: если файл не получается открыть
        """
        try:
            self.docx_file = docx.Document(self.file)
        except docx.opc.exceptions.PackageNotFoundError as e:
            raise ValueError(f"Файл имеет неверный формат (.docx) или поврежден")
        except Exception as e:
            raise ValueError(f"Файл не удаётся прочитать по техническим причинам: {e}")

    def read_document(self):
        """
        Чтение содержимого документа.
        
        :return: Экземпляр объекта Doc.
        """

        # вопрос, будем ли мы проверять как-либо ещё дополнительно форматирование
        # т.е. содержит ли документ нужные заделы, например
        # содержит ли нужные таблицы с данными по типу темы, автор и д.т.
        
        if isinstance(self.file, str):
            self._check_file_exists()
            self._check_file_format()
        else:
            self._load_file_from_buffer()

        if not hasattr(self, 'docx_file'):
            raise AttributeError("Файл не был открыт")

        doc = Doc()
        self.doc = doc

        self._read_structure(doc)
        self._read_info(doc)
        self._make_doc_structure_accurate(doc)

        return doc

    def _read_structure(self, doc: Doc):
        """
        Чтение основной структуры документа:
        проходит по всем параграфам начиная с параграфа "СОДЕРЖАНИЕ",
        заканчивая параграфом "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        сохраняет в виде вложенного списка в doc.structure,
        пропускает пустые параграфы (в изначальном документе - переносы строк)
        """

        found_content_section = False
        found_references_section = False
        current_section = DocSection()
        doc.structure.append(current_section)

        for paragraph in self.docx_file.paragraphs:
            text = paragraph.text
            # text = self._cut_code_paragraph(paragraph)
            text_norm = text.strip().lower()

            if len(text_norm) == 0:
                continue

            elif "содержание" == text_norm:
                found_content_section = True
                continue
            elif "список использованных источников" == text_norm:
                found_references_section = True
                break

            elif found_content_section and not found_references_section:
                level = self._get_heared_level(paragraph)
                
                if level == 0:
                    if len(current_section.text) == 0:
                        current_section.text = text
                    else:
                        current_section.text += ' ' + text

                elif current_section.name is None and len(current_section.text) == 0:
                    current_section.level = level
                    current_section.name = text

                elif current_section.level is None \
                    or level == current_section.level:
                    new_section = DocSection()
                    new_section.level = level
                    new_section.name = text
                    if current_section.upper is not None:
                        new_section.upper = current_section.upper
                        current_section = current_section.upper
                        current_section.structure.append(new_section)
                        current_section = new_section
                    else:
                        current_section = new_section
                        doc.structure.append(current_section)
                
                elif level > current_section.level:
                    new_section = DocSection()
                    new_section.upper = current_section
                    new_section.name = text
                    new_section.level = level
                    current_section.structure.append(new_section)
                    current_section = new_section

                elif level < current_section.level:
                    while current_section.upper is not None and level <= current_section.level:
                        current_section = current_section.upper
                    new_section = DocSection()
                    new_section.name = text
                    new_section.level = level
                    if current_section.level is None or \
                        current_section.upper is None and level <= current_section.level:
                        current_section = new_section
                        doc.structure.append(current_section)
                    else:
                        new_section.upper = current_section
                        current_section.structure.append(new_section)
                        current_section = new_section

    def _get_heared_level(self, paragraph):
        level = 0
        s = paragraph.style.name.lower().split()
        if len(s) > 1 and s[0] == 'heading':
            level = int(s[len(s) - 1])
        return level
    
    # пока не работает вообще, там криво читается, но читаться будет примерно таким образом,
    # просто надо переделать немного, чтобы с отальным кодом коннектилось
    # def _cut_code_paragraph(self, paragraph):
    #     text_no_courier = []
    #     for run in paragraph.runs:
    #         t = run.text.strip()
    #         if len(t) == 0 and \
    #             isinstance(run.font.name, str) and\
    #             not 'courier' in run.font.name.lower():
    #             text_no_courier.append(run.text.lower())
    #     return ' '.join(text_no_courier) if len(text_no_courier) >= 0 else ''
    
    def _read_info(self, doc: Doc):
        """
        Чтение основной информации о документе:
        doc.name; doc.author; doc.academic_supervisor; doc.year
        """
        try:
            try:
                theme = None
                t_i = None
                year = None
                y_i = None
                
                for i in range(len(self.docx_file.paragraphs)):
                    p = self.docx_file.paragraphs[i]
                    if t_i is not None \
                        and year is None \
                        and len(p.text.strip()) > 0:
                        if y_i is None:
                            y_i = i
                        else:
                            year = int(p.text.strip())
                            y_i = i
                            break
                    if theme is None and str(p.text.lower()).startswith('тема'):
                        t_i = i
                        theme = str(p.text[p.text.find(':') + 1:]).strip()
                doc.name = theme
                doc.year = year
            except:
                raise Exception("unable to read name and year")

            if len(self.docx_file.tables) < 2:
                raise Exception("error checking formating of docx file: len(self.docx_file.tables) < 2")
            if len(self.docx_file.tables[1].rows) < 3:
                raise Exception("error checking formating of docx file: len(self.docx_file.tables[1].rows) < 3")
            if len(self.docx_file.tables[1].columns) < 3:
                raise Exception("error checking formating of docx file: len(self.docx_file.tables[1].columns) < 3")
            if self.docx_file.tables[1].rows[0].cells[0].text != 'Студент':
                raise Exception("error checking formating of docx file: self.docx_file.tables[1].rows[0].cells[0].text != 'Студент'")
            if self.docx_file.tables[1].rows[2].cells[0].text != 'Руководитель':
                raise Exception("error checking formating of docx file: self.docx_file.tables[1].rows[2].cells[0].text != 'Руководитель'")
            
            try:
                h = 0
                author = ''
                i = len(self.docx_file.tables[1].rows[h].cells) - 1
                while i >= 0 and len(author) == 0:
                    author = self.docx_file.tables[1].rows[h].cells[i].text.strip()
                    i -= 1
                doc.author = author
    
                h = 2
                academic_supervisor = ''
                i = len(self.docx_file.tables[1].rows[h].cells) - 1
                while i >= 0 and len(academic_supervisor) == 0:
                    academic_supervisor = self.docx_file.tables[1].rows[h].cells[i].text.strip()
                    i -= 1
                doc.academic_supervisor = academic_supervisor
            except:
                raise Exception("formating check passed, but have a problem during reading author and academic_supervisor")
            
        except Exception as e:
            raise Exception(f"something wrong will reading info of documents (name, year, author, academic_supervisor): \n{e}")

    def _make_doc_structure_accurate(self, doc: Doc):
        i = 0
        b = False
        index_start = -1
        index_end = -1
        while i < len(doc.structure):
            el = doc.structure[i]
            if not b and 'введение' in str(el.name).lower().strip():
                b = True
                index_start = i
            if b and 'заключение' in str(el.name).lower().strip():
                b = False
                index_end = i
                break
            i += 1
        if index_start != -1 and index_end != -1:
            chapt = DocSection()
            chapt.name = 'Структура'
            i = index_start + 1
            while i < index_end:
                el = doc.structure[i]
                el.upper = None
                chapt.structure.append(el)
                doc.structure.pop(i)
                index_end -= 1
            doc.structure.insert(i, chapt)


