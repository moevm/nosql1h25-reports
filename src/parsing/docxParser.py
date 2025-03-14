import docx
import os
from typing import Union, BinaryIO
from io import BytesIO

from docClasses import Doc, DocSection


class DocxParser:
    def __init__(self, file: str | Union[BinaryIO, BytesIO]):
        self.file = file
        self.docx_file = None

    def _check_file_exists(self):
        """
        Проверяет, существует ли файл и является ли он строкой.
        
        :raises ValueError: если путь к файлу не валиден
        :raises FileNotFoundError: если файл не существует
        """
        if not self.file or not isinstance(self.file, str):
            raise ValueError("Путь к файлу должен быть строкой и не пустым")
        try:
            if not os.path.exists(self.file) or not os.path.isfile(self.file):
                raise FileNotFoundError(f"Файл '{self.file}' не найден")
        except ImportError as e:
            raise("Возможная ошибка: ", str(e))

    def _check_file_format(self):
        """
        Проверяет, является ли файл форматом .docx.
        
        :raises ValueError: если файл имеет неверный формат
        """
        try:
            self.docx_file = docx.Document(self.file)
        except docx.exc.InvalidFileException as e:
            raise ValueError(f"Файл '{self.file}' имеет неверный формат (.docx)")
    
    def _load_life_from_buffer(self):
        """
        Предполагается, что на вход подан какой-то буффер или массив байтов.
        
        :raises ValueError: если файл не получается открыть
        """
        try:
            self.docx_file = docx.Document(self.file)
        except docx.exc.InvalidFileException as e:
            raise ValueError(f"Файл '{self.file}' имеет неверный формат (.docx)")
        except e:
            raise ValueError(f"Файл не удаётся прочитать по техническим причинам")

    def read_document(self):
        """
        Чтение содержимого документа.
        
        :return: Экземпляр объекта Doc.
        """

        # вопрос, будем ли мы проверять как-либо ещё дополнительно форматирование
        # т.е. содержит ли документ нужные заделы, например
        # содержит ли нужные таблицы с данными по типу темы, автор и д.т.
        
        if isinstance(self.file, str):
            self._check_file_exists()
            self._check_file_format()
        else:
            self._load_life_from_buffer()

        if not hasattr(self, 'docx_file'):
            raise AttributeError("Файл не был открыт")

        doc = Doc()

        self._read_structure(doc)
        self._read_info(doc)

        return doc

    def _read_structure(self, doc: Doc):
        """
        Чтение основной структуры документа:
        проходит по всем параграфам начиная с параграфа "СОДЕРЖАНИЕ",
        заканчивая параграфом "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        сохраняет в виде вложенного списка в doc.structure,
        пропускает пустые параграфы (в изначальном документе - переносы строк)
        """

        found_content_section = False
        found_references_section = False
        current_section = DocSection()
        doc.structure.append(current_section)

        for paragraph in self.docx_file.paragraphs:
            text = paragraph.text.strip().lower()
            # text = self._cut_code_paragraph(paragraph)

            if len(text) == 0:
                continue

            elif "содержание" == text:
                found_content_section = True
                continue
            elif "список использованных источников" == text:
                found_references_section = True
                break

            elif found_content_section and not found_references_section:
                level = self._get_heared_level(paragraph)
                
                if level == 0:
                    if len(current_section.text) == 0:
                        current_section.text = text
                    else:
                        current_section.text += ' ' + text

                elif current_section.title is None and len(current_section.text) == 0:
                    current_section.level = level
                    current_section.title = text

                elif current_section.level is None \
                    or level == current_section.level:
                    new_section = DocSection()
                    new_section.level = level
                    new_section.title = text
                    if current_section.upper is not None:
                        new_section.upper = current_section.upper
                        current_section = current_section.upper
                        current_section.structure.append(new_section)
                        current_section = new_section
                    else:
                        current_section = new_section
                        doc.structure.append(current_section)
                
                elif level > current_section.level:
                    new_section = DocSection()
                    new_section.upper = current_section
                    new_section.title = text
                    new_section.level = level
                    current_section.structure.append(new_section)
                    current_section = new_section

                elif level < current_section.level:
                    while current_section.upper is not None and level <= current_section.level:
                        current_section = current_section.upper
                    new_section = DocSection()
                    new_section.title = text
                    new_section.level = level
                    if current_section.level is None or \
                        current_section.upper is None and level <= current_section.level:
                        current_section = new_section
                        doc.structure.append(current_section)
                    else:
                        new_section.upper = current_section
                        current_section.structure.append(new_section)
                        current_section = new_section

    def _get_heared_level(self, paragraph):
        level = 0
        s = paragraph.style.name.lower().split()
        if len(s) > 1 and s[0] == 'heading':
            level = int(s[len(s) - 1])
        return level
    
    # пока не работает вообще, там криво читается, но читаться будет примерно таким образом,
    # просто надо переделать немного, чтобы с отальным кодом коннектилось
    # def _cut_code_paragraph(self, paragraph):
    #     text_no_courier = []
    #     for run in paragraph.runs:
    #         t = run.text.strip()
    #         if len(t) == 0 and \
    #             isinstance(run.font.name, str) and\
    #             not 'courier' in run.font.name.lower():
    #             text_no_courier.append(run.text.lower())
    #     return ' '.join(text_no_courier) if len(text_no_courier) >= 0 else ''
    
    def _read_info(self, doc: Doc):
        """
        Чтение основной информации о документе:
        doc.title; doc.author; doc.sci_director; doc.count_wolds; doc.count_symbols
        """
        pass